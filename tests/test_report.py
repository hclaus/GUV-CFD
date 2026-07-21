import json
import time

import pytest

from guvcfd.report import generate_report_docx, _format_elapsed, _run_timing, _trust_status_rows


@pytest.fixture(autouse=True)
def _fast_system_info(monkeypatch):
    # get_system_info() shells out to PowerShell/WMI (~1-2s, and not
    # available on non-Windows CI) - every report test goes through
    # generate_report_docx, so fake it out once here rather than eating
    # that cost (and that dependency) in every single test.
    monkeypatch.setattr(
        "guvcfd.report.get_system_info",
        lambda: {"cpu": "Test CPU Model", "ram_gb": 16.0, "gpu": "Test GPU"},
    )


def test_missing_run_settings_raises_clear_error(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "results.json").write_text("{}")
    with pytest.raises(FileNotFoundError, match="run a full simulation"):
        generate_report_docx(case_dir, str(tmp_path / "out.docx"))


def test_missing_results_raises_clear_error(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text("{}")
    with pytest.raises(FileNotFoundError, match="run a full simulation"):
        generate_report_docx(case_dir, str(tmp_path / "out.docx"))


def test_missing_guv_path_raises_clear_error(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps({"ach": 3.0}))  # no guv_path
    (tmp_path / "results.json").write_text("{}")
    with pytest.raises(FileNotFoundError, match="predates report support"):
        generate_report_docx(case_dir, str(tmp_path / "out.docx"))


_REAL_SETTINGS = {
    "ach": 3.0, "z-value": 2.0,
    "inlet-wall": "xMin", "inlet-y-input": 1.5, "inlet-z-input": 2.295,
    "inlet-size-w": 0.3, "inlet-size-h": 0.3,
    "outlet-wall": "xMax", "outlet-y-input": 1.5, "outlet-z-input": 0.405,
    "outlet-size-w": 0.3, "outlet-size-h": 0.3,
    "fan-enable": False,
    "guv_path": r"c:\Users\hukcl\Documents\Python\Illuminator2\illuminate-v4\4x3x2.7.guv",
}

_STEADY_STATE_RESULTS = {
    "target_T_ss": 0.3,
    "phase1": {"T_ss": 0.2548, "converged": True, "iterations": 8000,
               "decay_curve": {"t": [0, 100, 200], "T": [0.0, 0.15, 0.2548]}},
    "phase2": {"T_ss": 0.0644, "converged": False, "iterations": 3000,
               "decay_curve": {"t": [0, 100, 200], "T": [0.2548, 0.12, 0.0644]}},
    "reduction_pct": 74.7,
    "eACH_uv_steady_state": 17.73,
    "eACH_uv_well_mixed": 15.2,
    "fluence_mean": 12.34,
}


def test_steady_state_report_does_not_crash_on_decay_only_fields(tmp_path):
    # Regression test: steady-state results.json has a totally different
    # schema (phase1/phase2/reduction_pct/eACH_uv_steady_state) than decay's
    # (ventilation_ach/eACH_uv_effective/decay_curve) - generate_report_docx
    # must dispatch on scenario type instead of assuming decay's fields.
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            all_text += "\n" + "\t".join(c.text for c in row.cells)
    assert "74.7%" in all_text
    assert "15.2" in all_text  # eACH_uv_well_mixed ("Calculated eACH" row)
    assert "17.73" in all_text  # eACH_uv_steady_state ("Simple CFD measured eACHCFD_s" row)
    assert "12.34" in all_text  # average fluence rate
    assert len(doc.inline_shapes) == 2  # room preview + phase-timeline curve pictures embedded


def test_steady_state_report_shows_windowed_stats_when_present(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["phase1"] = dict(results["phase1"], T_ss_cv=0.012, T_ss_window_frac=0.15)
    results["phase2"] = dict(results["phase2"], T_ss_cv=0.14, T_ss_window_span=456)
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    assert "Steady state T, calculated from  moving average (15% of last results)" in text
    assert "CV (15% of last results)" in text
    assert "1.2%" in text
    assert "Steady State TSS2, calculated from moving average (last 456 iterations)" in text
    assert "CV (last 456 iterations)" in text
    assert "14.0%" in text


def test_steady_state_report_uses_generic_window_phrase_when_window_fields_absent(tmp_path):
    # Older results.json predating live-volAverage tracking has no
    # T_ss_window_frac/_span - the label must still read sensibly (not
    # leave the template's own literal "<...>" placeholder in the output).
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    assert "a trailing window of results" in text  # phase 1 fallback phrase
    assert "a trailing window" in text  # phase 2 fallback phrase
    assert "0.2548" in text  # phase 1 T_ss still shown
    assert "0.0644" in text  # phase 2 T_ss still shown


def test_decay_report_shows_fluence_mean(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    decay_results = {
        "ventilation_ach": 3.0, "eACH_uv_well_mixed": 10.27, "eACH_uv_effective": 8.97,
        "mixing_efficiency": 0.873, "total_ach_effective": 11.97,
        "decay_curve": {"t_seconds": [0, 10], "volAverage_T": [1.0, 0.9]},
        "fluence_mean": 5.678,
    }
    (tmp_path / "results.json").write_text(json.dumps(decay_results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    all_text = ""
    for table in doc.tables:
        for row in table.rows:
            all_text += "\n" + "\t".join(c.text for c in row.cells)
    assert "5.678" in all_text


def test_steady_state_report_shows_corrected_fields_when_present(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["ventilation_ach_measured"] = 2.55
    results["eACH_uv_steady_state_corrected"] = 18.1
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    all_text = ""
    for table in doc.tables:
        for row in table.rows:
            all_text += "\n" + "\t".join(c.text for c in row.cells)
    assert "2.55" in all_text
    assert "18.1" in all_text


def test_steady_state_report_recomputes_injection_rate_for_old_runs(tmp_path):
    # injection_rate_total predates the field in old results.json files -
    # it's deterministic from room volume/ACH/target_T_ss, so the report
    # should show the real number instead of "n/a" for a case dir that
    # predates the field being saved.
    from guvcfd.contaminant_source import compute_source_strength
    from guv_calcs import Project

    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    assert "injection_rate_total" not in _STEADY_STATE_RESULTS
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    room = next(iter(Project.load(_REAL_SETTINGS["guv_path"]).rooms.values()))
    expected_G = compute_source_strength(room.x * room.y * room.z, _REAL_SETTINGS["ach"],
                                          _STEADY_STATE_RESULTS["target_T_ss"])

    from docx import Document
    doc = Document(out_path)
    all_text = ""
    for table in doc.tables:
        for row in table.rows:
            all_text += "\n" + "\t".join(c.text for c in row.cells)
    assert f"{expected_G:.4g}" in all_text
    injection_row = next(row for row in all_text.split("\n") if "Source injection rate" in row)
    assert "n/a" not in injection_row  # the real number, not the "no injection_rate_total" fallback


def test_steady_state_report_recomputes_theoretical_each_uv_for_old_runs(tmp_path):
    # eACH_uv_well_mixed predates app.py copying it out of setup_case()'s
    # summary into results.json - it's exactly Z * fluence_mean * 3.6 (see
    # fluence.py), so an older results.json missing it should still show
    # the real number instead of "n/a".
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    del results["eACH_uv_well_mixed"]
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    expected = _REAL_SETTINGS["z-value"] * results["fluence_mean"] * 3.6

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    theoretical_row = next(row for row in text.split("\n") if row.startswith("Calculated eACH"))
    assert f"{expected:.4g}" in theoretical_row


def test_report_always_includes_t_field_note(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "T is the OpenFOAM field name" in all_text


def test_report_flags_non_uniform_mixing_when_monitoring_points_diverge(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["monitoring"] = {
        "Patient": {
            "phase1": {"volAverage_T": [0.0, 0.1957]},
            "phase2": {"volAverage_T": [0.1957, 0.0122]},
        },
    }
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "NOT well mixed" in all_text
    assert "Patient" in all_text


def test_format_elapsed():
    assert _format_elapsed(45) == "45s"
    assert _format_elapsed(125) == "2m 5s"
    assert _format_elapsed(5445) == "1h 30m 45s"


def test_run_timing_prefers_recorded_fields(tmp_path):
    results = {"run_started_at": "2026-07-13T14:30:00", "run_elapsed_seconds": 90}
    started_at, elapsed = _run_timing(str(tmp_path), results)
    assert started_at.isoformat() == "2026-07-13T14:30:00"
    assert elapsed == 90


def test_run_timing_falls_back_to_file_mtimes(tmp_path):
    settings_path = tmp_path / "run_settings.json"
    results_path = tmp_path / "results.json"
    settings_path.write_text("{}")
    time.sleep(0.05)
    results_path.write_text("{}")

    started_at, elapsed = _run_timing(str(tmp_path), {})
    assert started_at is not None
    assert elapsed >= 0


def test_run_timing_returns_none_when_nothing_available(tmp_path):
    started_at, elapsed = _run_timing(str(tmp_path), {})
    assert started_at is None
    assert elapsed is None


def _table_text(doc):
    text = ""
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + "\t".join(c.text for c in row.cells)
    return text


def _results_table(doc):
    # The results-table template has a fixed 24 rows - doc.tables[0]/[-1]
    # aren't reliable once other sections (metadata, Room Setup, optional
    # Monitoring Results) are relocated around it in the final document.
    return next(t for t in doc.tables if len(t.rows) == 24)


def test_report_shows_recorded_simulation_date_and_elapsed_time(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["run_started_at"] = "2026-07-13T14:30:00"
    results["run_elapsed_seconds"] = 5445  # 1h 30m 45s
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    assert "2026-07-13 14:30" in text
    assert "1h 30m 45s" in text


def test_report_falls_back_to_file_times_when_run_timing_not_recorded(tmp_path):
    case_dir = str(tmp_path)
    settings_path = tmp_path / "run_settings.json"
    results_path = tmp_path / "results.json"
    settings_path.write_text(json.dumps(_REAL_SETTINGS))
    assert "run_started_at" not in _STEADY_STATE_RESULTS
    time.sleep(0.05)
    results_path.write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    assert "Simulation date" in text
    assert "Total elapsed time" in text


def test_report_shows_system_info(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    assert "Test CPU Model" in text
    assert "16.0 GB" in text
    assert "Test GPU" in text
    assert "not used" in text


def test_report_header_paragraphs_show_guv_path_settings_path_and_case_dir(tmp_path):
    case_dir = str(tmp_path)
    settings = dict(_REAL_SETTINGS)
    settings["settings_path"] = r"c:\Users\hukcl\Documents\Python\GUV-CFD\projects\my_room.guvcfd"
    (tmp_path / "run_settings.json").write_text(json.dumps(settings))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert f"Illuminate room design file: {_REAL_SETTINGS['guv_path']}" in all_text
    assert "CFD Project file: c:\\Users\\hukcl\\Documents\\Python\\GUV-CFD\\projects\\my_room.guvcfd" in all_text
    assert f"OpenFoam directory: {case_dir}" in all_text


def test_report_header_shows_na_for_missing_settings_path(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))  # no settings_path key
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "CFD Project file: n/a" in all_text


def test_room_setup_shows_z_units_and_second_opening_and_injection_and_monitoring_points(tmp_path):
    case_dir = str(tmp_path)
    settings = dict(_REAL_SETTINGS)
    settings.update({
        "inlet2-enable": True, "inlet2-wall": "ceiling",
        "inlet2-y-input": 2.0, "inlet2-z-input": 1.5,
        "inlet2-size-w": 0.3, "inlet2-size-h": 0.3,
        "outlet2-enable": False,
        "source_center": [2.0, 1.5, 1.5],
        "monitoring_points": [{"name": "Patient", "x": 1.0, "y": 1.0, "z": 1.2, "cells_per_side": 4}],
    })
    (tmp_path / "run_settings.json").write_text(json.dumps(settings))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    assert "cm²/mJ" in text  # Z units fixed, not the old m^2/J
    assert "Inlet 2" in text and "ceiling" in text
    assert "Outlet 2" not in text  # disabled - omitted entirely
    assert "Injection point" in text
    assert "Monitoring point: Patient" in text


def test_steady_state_results_table_has_calculated_and_total_ach_rows(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["ventilation_ach_measured"] = 2.55
    results["eACH_uv_steady_state_corrected"] = 18.1
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    assert "Calculated eACH" in text
    assert "15.2" in text  # eACH_uv_well_mixed
    assert "Effective pathogen (mechanical) ACHeff" in text
    assert "2.55" in text
    assert "Room ventilation pathogen removal efficacy EACHeff" in text
    assert f"{2.55 / _REAL_SETTINGS['ach'] * 100:.1f}%" in text
    assert "True CFD measured eACHCFD" in text
    assert "18.1" in text
    assert "Simple CFD measured eACHCFD_s" in text
    assert "17.73" in text  # eACH_uv_steady_state
    total_row = next(row for row in text.split("\n") if row.startswith("Total ACH in room"))
    assert f"{2.55 + 18.1:.4g}" in total_row


def test_steady_state_results_table_total_ach_falls_back_to_na(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))  # no *_measured fields
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    total_row = next(row for row in text.split("\n") if row.startswith("Total ACH in room"))
    assert total_row.strip().endswith("n/a")


def test_steady_state_results_table_uvgi_effectiveness_rows_match(tmp_path):
    # "True" and "Simple" UVGI Effectiveness are algebraically forced to be
    # the same number (both reduce to 1 - T_ss2/T_ss1 regardless of which
    # ACH basis feeds them) - see the chat discussion when this format was
    # approved. Must show the same value in both rows, not a fabricated
    # "True" figure.
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    table = _results_table(doc)
    assert table.rows[22].cells[1].text == "74.7%"
    assert table.rows[23].cells[1].text == "74.7%"


def test_steady_state_results_table_achieff_label_notes_ach_source(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["ventilation_ach_measured"] = 2.55
    results["eACH_uv_steady_state_corrected"] = 18.1
    results["ach_source"] = "extrapolated_T_infinity"
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    table = _results_table(doc)
    assert table.rows[16].cells[0].text == "Effective pathogen (mechanical) ACHeff (using extrapolated T∞)"


def test_steady_state_results_table_extrapolation_row_shows_na_when_unavailable(tmp_path):
    # Phase 1/2 in _STEADY_STATE_RESULTS have no T_inf_extrapolated (the
    # fit never ran/converged) - row 9/13 must say so explicitly rather
    # than showing a stale or blank cell.
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    table = _results_table(doc)
    assert table.rows[9].cells[1].text == "n/a (extrapolation unavailable)"
    assert table.rows[13].cells[1].text == "n/a (extrapolation unavailable)"


def test_steady_state_results_table_footnotes_are_filled_with_correct_style(tmp_path):
    # The filled-in footnote text must land in a normal-styled run, not
    # the small-superscript run that carries the footnote's own auto-number
    # marker (w:footnoteRef) - overwriting that run's text was the original
    # font-size bug caught during review of the manually-generated sample.
    import zipfile
    from lxml import etree

    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["phase1"] = dict(results["phase1"], live={"t": list(range(8000)), "T": [0.0] * 8000})
    results["phase2"] = dict(results["phase2"], live={"t": list(range(4000)), "T": [0.0] * 4000})
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": W}
    with zipfile.ZipFile(out_path) as z:
        footnotes_xml = etree.fromstring(z.read("word/footnotes.xml"))

    for fid, expect_n in (("4", 4000), ("5", 2000)):
        footnote = next(fn for fn in footnotes_xml.findall("w:footnote", ns)
                         if fn.get(f"{{{W}}}id") == fid)
        runs = footnote.findall(".//w:r", ns)
        marker_run = next(r for r in runs if r.find(f"{{{W}}}footnoteRef") is not None)
        assert marker_run.find(f"{{{W}}}t") is None  # marker run's text untouched
        text_run = next(r for r in runs if r.find(f"{{{W}}}footnoteRef") is None
                         and r.find(f"{{{W}}}t") is not None and r.find(f"{{{W}}}t").text)
        style = text_run.find(f".//{{{W}}}rStyle")
        assert style is None  # normal style, not FootnoteReference (small/superscript)
        assert f"{expect_n} simulation points" in text_run.find(f"{{{W}}}t").text


def test_monitoring_results_heading_renamed(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["monitoring"] = {
        "Patient": {
            "phase1": {"volAverage_T": [0.0, 0.1957]},
            "phase2": {"volAverage_T": [0.1957, 0.0122]},
        },
    }
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Monitoring Results" in headings
    assert "Monitoring Locations" not in headings


def test_monitoring_row_shows_windowed_t_ss_and_cv_when_present(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["monitoring"] = {
        "Patient": {
            "phase1": {"volAverage_T": [0.0, 0.21], "T_ss": 0.207, "T_ss_cv": 0.052},
            "phase2": {"volAverage_T": [0.21, 0.003], "T_ss": 0.0020, "T_ss_cv": 0.541},
        },
    }
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    text = _table_text(doc)
    assert "T_ss1=0.207" in text
    assert "T_ss2=0.002" in text
    assert "CV1=5.2%" in text
    assert "CV2=54.1%" in text


def test_steady_state_report_embeds_phase_timeline_curve(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    (tmp_path / "results.json").write_text(json.dumps(_STEADY_STATE_RESULTS))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    assert len(doc.inline_shapes) == 2  # case-setup preview + phase-timeline curve


def test_steady_state_report_skips_curve_picture_when_no_decay_curve_data(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    results = dict(_STEADY_STATE_RESULTS)
    results["phase1"] = {"T_ss": 0.2548, "converged": True, "iterations": 8000}  # no decay_curve
    results["phase2"] = {"T_ss": 0.0644, "converged": False, "iterations": 3000}
    (tmp_path / "results.json").write_text(json.dumps(results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    assert len(doc.inline_shapes) == 1  # only the case-setup preview


def test_decay_report_embeds_decay_curve_picture(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    decay_results = {
        "ventilation_ach": 3.0, "eACH_uv_well_mixed": 10.27, "eACH_uv_effective": 8.97,
        "mixing_efficiency": 0.873, "total_ach_effective": 11.97,
        "decay_curve": {"t_seconds": [0, 10], "volAverage_T": [1.0, 0.9]},
        "fluence_mean": 5.678,
    }
    (tmp_path / "results.json").write_text(json.dumps(decay_results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    assert len(doc.inline_shapes) == 2  # case-setup preview + decay-curve picture


def test_trust_status_rows_flow_converged_and_ach_ok():
    rows = _trust_status_rows({
        "flow_converged": True,
        "ach_delivery": {"measured_ach": 1.48, "nominal_ach": 1.5, "ratio": 0.987, "within_tolerance": True},
    })
    labels = dict(rows)
    assert "Converged" in labels["Flow field"]
    assert "OK" in labels["Ventilation delivery"]


def test_trust_status_rows_flow_accepted_via_oscillation():
    rows = _trust_status_rows({"flow_converged": False, "ach_delivery": None})
    labels = dict(rows)
    assert "Accepted via bounded oscillation" in labels["Flow field"]
    assert labels["Ventilation delivery"] == "Not checked (older run)"


def test_trust_status_rows_flags_ach_mismatch():
    rows = _trust_status_rows({
        "flow_converged": True,
        "ach_delivery": {"measured_ach": 0.57, "nominal_ach": 1.5, "ratio": 0.38, "within_tolerance": False},
    })
    labels = dict(rows)
    assert "MISMATCH" in labels["Ventilation delivery"]


def test_trust_status_rows_missing_fields_report_not_available():
    rows = _trust_status_rows({})
    labels = dict(rows)
    assert labels["Flow field"] == "Not available (older run, or flow convergence was skipped)"
    assert labels["Ventilation delivery"] == "Not checked (older run)"
    assert "Phase 1 mass balance" not in labels


def test_trust_status_rows_includes_phase1_mass_balance_when_present():
    rows = _trust_status_rows({
        "phase1": {"mass_balance": {"ratio": 0.73, "within_tolerance": False}},
    })
    labels = dict(rows)
    assert "NOT balanced" in labels["Phase 1 mass balance"]


def test_trust_status_rows_reported_in_generated_docx(tmp_path):
    case_dir = str(tmp_path)
    (tmp_path / "run_settings.json").write_text(json.dumps(_REAL_SETTINGS))
    decay_results = {
        "ventilation_ach": 3.0, "eACH_uv_well_mixed": 10.27, "eACH_uv_effective": 8.97,
        "mixing_efficiency": 0.873, "total_ach_effective": 11.97,
        "decay_curve": {"t_seconds": [0, 10], "volAverage_T": [1.0, 0.9]},
        "fluence_mean": 5.678,
        "flow_converged": False,
        "ach_delivery": {"measured_ach": 0.57, "nominal_ach": 1.5, "ratio": 0.38, "within_tolerance": False},
    }
    (tmp_path / "results.json").write_text(json.dumps(decay_results))
    out_path = str(tmp_path / "out.docx")

    generate_report_docx(case_dir, out_path)

    from docx import Document
    doc = Document(out_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Convergence & Trust" in full_text
