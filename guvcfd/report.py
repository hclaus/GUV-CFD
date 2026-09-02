"""Generate a one-page-ish .docx summary of a completed decay run: room
setup parameters, a rendered picture of the case setup (inlet/outlet/fan/
lamps), and the key result numbers - for sharing outside the GUI itself.
"""
import json
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree
from guv_calcs import Project

from .contaminant_source import compute_source_strength
from .decay_analysis import windowed_stats_detrended, mechanical_mixing_efficiency_pct as _mech_mixing_eff_pct, migrate_result_keys
from .monitoring_points import mixing_uniformity_note, point_reduction_basis
from .result_figures import decay_figure, steady_state_figure
from .system_info import get_system_info
from .visualization import center_frac_for_wall, plot_case


def _format_elapsed(seconds):
    seconds = max(0, int(seconds))
    h, rem = divmod(seconds, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h}h {m}m {s}s"
    if m:
        return f"{m}m {s}s"
    return f"{s}s"


def _run_timing(case_dir, results):
    """(started_at, elapsed_seconds) for the "Simulation date"/"Total
    elapsed time" report rows - from results.json's own run_started_at/
    run_elapsed_seconds if this run recorded them, else a rough fallback
    from run_settings.json's/results.json's file-modified times (written
    near the start and end of a run respectively) for older or
    still-in-progress case directories that predate that tracking.
    """
    if results.get("run_started_at") is not None:
        started_at = datetime.fromisoformat(results["run_started_at"])
        return started_at, results.get("run_elapsed_seconds")
    try:
        start_mtime = (Path(case_dir) / "run_settings.json").stat().st_mtime
        end_mtime = (Path(case_dir) / "results.json").stat().st_mtime
        return datetime.fromtimestamp(start_mtime), max(0, end_mtime - start_mtime)
    except OSError:
        return None, None

# What "T" actually is - shown once under the Results heading in both the
# .docx report and the Analysis tab (imported from here by app.py) since
# neither this simulation nor OpenFOAM itself assigns it a physical unit.
T_FIELD_NOTE = (
    "Note: T is the OpenFOAM field name for the transported scalar this "
    "whole simulation tracks - the substance being reduced, per unit volume. "
    "In a GUV disinfection context this is typically a pathogen "
    "concentration, e.g. CFU/m³ (colony-forming units per cubic meter) or "
    "an equivalent airborne-contaminant unit; the CFD itself is "
    "unit-agnostic and just tracks relative concentration."
)

# Shown alongside the "measured/effective ventilation ACH" rows below so the
# number isn't mistaken for a claim that less air is being delivered - the
# inlet flow rate is fixed at the nominal ACH by the boundary condition
# itself (mass conservation guarantees it all exits through the outlet).
# This metric instead reflects how well that air mixes into the room.
# Two different ways to measure it exist (see
# steady_state_pipeline.run_steady_state_scenario's "ventilation_measurement_method"
# field) - _effective_ach_note picks the right explanation for whichever
# one actually produced this result.
EFFECTIVE_ACH_NOTE = (
    "Note: the effective/measured ventilation ACH above is not the airflow "
    "rate delivered by the inlet - that's fixed at the nominal design ACH "
    "by the boundary condition. It's the well-mixed-equivalent rate implied "
    "by Phase 1's room-average steady-state concentration, and reads lower "
    "than nominal when the room mixes imperfectly (e.g. inlet/outlet "
    "short-circuiting on the same wall)."
)

EFFECTIVE_ACH_NOTE_CONTROL = (
    "Note: the effective/measured ventilation ACH above is not the airflow "
    "rate delivered by the inlet - that's fixed at the nominal design ACH "
    "by the boundary condition. It's measured directly by a dedicated UV-off "
    "control run (uniform initial concentration, ventilation only, no source) "
    "run alongside the main scenario - see the OpenFOAM Notes help page for "
    "why this is preferred over deriving it from Phase 1's own point-source "
    "buildup. Reads lower than nominal when the room mixes imperfectly (e.g. "
    "inlet/outlet short-circuiting on the same wall)."
)


def _effective_ach_note(results):
    """Pick EFFECTIVE_ACH_NOTE vs EFFECTIVE_ACH_NOTE_CONTROL based on which
    method actually measured ventilation_ach_measured for this result -
    older results.json files without "ventilation_measurement_method" fall
    back to the Phase-1-based wording (the only method that existed then).
    """
    if results.get("ventilation_measurement_method") == "control_run":
        return EFFECTIVE_ACH_NOTE_CONTROL
    return EFFECTIVE_ACH_NOTE

_ROW_LABELS_ROOM = [
    ("Room dimensions", lambda r, s: f"{r.x:.3g} x {r.y:.3g} x {r.z:.3g} {r.units}"),
    ("Lamps", lambda r, s: str(len(r.lamps))),
    ("Ventilation ACH", lambda r, s: f"{s['ach']:.3g} /hr"),
    ("UV inactivation constant Z", lambda r, s: f"{s['z-value']:.3g} cm²/mJ"),
    ("Inlet", lambda r, s: f"{s['inlet-wall']}, y={s['inlet-y-input']:.3g}m "
                           f"z={s['inlet-z-input']:.3g}m, size={s['inlet-size-w']:.3g}x"
                           f"{s['inlet-size-h']:.3g}m"),
    ("Outlet", lambda r, s: f"{s['outlet-wall']}, y={s['outlet-y-input']:.3g}m "
                            f"z={s['outlet-z-input']:.3g}m, size={s['outlet-size-w']:.3g}x"
                            f"{s['outlet-size-h']:.3g}m"),
]

_ROW_LABELS_FAN = [
    ("Mixing fan", lambda s: f"{s['fan-speed']:.3g} m/s, direction={s['fan-direction']}, "
                             f"position=({s['fan-x-input']:.3g}, {s['fan-y-input']:.3g}, "
                             f"{s['fan-z-input']:.3g})m, radius={s['fan-radius']:.3g}m"),
]

# Decay reports are seeded from their own approved template (a fixed
# 17-row results table plus a 2-row per-point monitoring table), the decay-
# mode counterpart to _RESULTS_TABLE_TEMPLATE_PATH below - see
# _decay_results_table_cell_values for the row layout.
_DECAY_RESULTS_TABLE_TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "decay_results_table_template.docx"


def _decay_reduction_ratio(eACH, ach):
    """Analytical steady-state pathogen reduction implied by a given
    ACH/eACH pair, WITHOUT running the (expensive) steady-state pipeline:
    for a fixed continuous source, T_ss ∝ 1/(ACH+eACH) (see
    compute_source_strength), so reduction = 1 - T_ss(with UV)/T_ss(no UV)
    = 1 - ACH/(ACH+eACH) = eACH/(ACH+eACH) - algebraically the same
    quantity steady_state_pipeline.compute_corrected_eACH_uv's own
    reduction_pct computes from an actual T_ss1/T_ss2 CFD ratio.

    An earlier version of this used a transient "1 - exp(-rate*1hr)"
    formula instead - correct for "how much decays away in one hour
    starting from a pulse," but not the steady-state (continuous-source)
    reduction this report actually needs; replaced after a real case
    showed a physically meaningless ~100% (saturated exp(-rate) with a
    large rate) where the true steady-state figure was ~11.9%.
    """
    return eACH / (ach + eACH) if (ach + eACH) else None


def combo_summary_metrics(detail):
    """The 5 headline numbers shown per combination on the Run Simulations
    tab's "Simulation results" columns and in the sweep-summary CSV
    (scenario_runs.write_sweep_summary_csv) - one place computing them so
    both stay consistent, for either a steady-state or decay-mode trimmed
    result dict (see scenario_runs._trim_report/_trim_decay_report).

    Returns a dict with keys total_reduction_pct, ach_efficiency_pct,
    uv_efficiency_pct, mechanical_mixing_efficiency_pct, est_ach_per_hr,
    est_each_per_hr, ach_t_measured_per_hr - any value that isn't
    computable from what's in `detail` (older results.json files, a
    control run that wasn't used, etc.) is None, not an error.

    ach_efficiency_pct: ach_delivery.ratio*100 - the measured/nominal
    ventilation-delivery ratio (same field either sim type already
    computes via run_pipeline.check_ach_delivery), i.e. how much of the
    nominal ACH the room's actual flow field delivers.

    ach_t_measured_per_hr: the shared UV-off control run's own measured
    ventilation rate (detail["ventilation_ach_measured"], or decay mode's
    older "ventilation_ach" name) - a scalar-decay (T) curve fit, a
    genuinely different measurement method from est_ach_per_hr above (a
    mesh flow-rate integration via check_ach_delivery, no scalar solve
    involved at all) - the two can legitimately disagree and both are
    worth seeing side by side, not just one.

    uv_efficiency_pct: the room's real (imperfectly-mixed) eACH_uv versus
    the idealized well-mixed prediction for the same Z - decay mode
    already computes this directly (mixing_efficiency[_corrected] via
    decay_analysis.write_results_summary); steady-state doesn't store the
    ratio itself, only eACH_uv_well_mixed, so it's computed here instead.

    mechanical_mixing_efficiency_pct: the UV-independent counterpart -
    measured effective ventilation removal vs. the measured DELIVERED
    flow rate (see decay_analysis.mechanical_mixing_efficiency_pct's own
    docstring for why this is a genuinely different question from either
    of the two above). Both sim types already store this field directly
    (see write_results_summary/app._finish_steady_state/scenario_runs.
    _run_scenario) - only recomputed here as a fallback for an
    older results.json predating this field.
    """
    is_steady_state = "reduction_pct" in detail or "reduction_pct_corrected" in detail
    ach_delivery = detail.get("ach_delivery") or {}
    ach_efficiency_pct = ach_delivery.get("ratio") * 100 if ach_delivery.get("ratio") is not None else None
    est_ach_per_hr = ach_delivery.get("measured_ach")
    # The shared UV-off control run's own measured ventilation rate - a
    # scalar-decay (T) fit, genuinely different from est_ach_per_hr above
    # (a mesh flow-rate integration via check_ach_delivery) - the two can
    # legitimately disagree, so both are exposed rather than picking one.
    # None whenever no control run was used (an older results.json, or a
    # sealed/no-ventilation combo).
    ach_t_measured_per_hr = detail.get("ventilation_ach_measured", detail.get("ventilation_ach"))

    if is_steady_state:
        total_reduction_pct = detail.get("reduction_pct_corrected", detail.get("reduction_pct"))
        est_each_per_hr = detail.get("eACH_uv_steady_state_corrected", detail.get("eACH_uv_steady_state"))
        well_mixed = detail.get("eACH_uv_well_mixed")
        uv_efficiency_pct = (est_each_per_hr / well_mixed * 100
                              if est_each_per_hr is not None and well_mixed else None)
    else:
        est_each_per_hr = detail.get("eACH_uv_actual", detail.get("eACH_uv_assuming_well_mixed"))
        total_reduction_pct = (_decay_reduction_ratio(est_each_per_hr, ach_t_measured_per_hr) * 100
                                if est_each_per_hr is not None and ach_t_measured_per_hr is not None else None)
        mixing_eff = detail.get("mixing_efficiency_actual", detail.get("mixing_efficiency"))
        uv_efficiency_pct = mixing_eff * 100 if mixing_eff is not None else None

    mechanical_mixing_pct = detail.get("mechanical_mixing_efficiency_pct")
    if mechanical_mixing_pct is None:
        mechanical_mixing_pct = _mech_mixing_eff_pct(detail)

    return {
        "total_reduction_pct": total_reduction_pct,
        "ach_efficiency_pct": ach_efficiency_pct,
        "uv_efficiency_pct": uv_efficiency_pct,
        "mechanical_mixing_efficiency_pct": mechanical_mixing_pct,
        "est_ach_per_hr": est_ach_per_hr,
        "est_each_per_hr": est_each_per_hr,
        "ach_t_measured_per_hr": ach_t_measured_per_hr,
    }


def _ci_suffix(ci95):
    """' (95% CI: lo–hi /hr)' for a (lo, hi) tuple in /hr, else ''.

    The CI is this fit's own regression uncertainty only (see
    compute_effective_eACH's docstring) - a real, standard OLS confidence
    interval a scientist reading this report would recognize, not a
    bespoke statistic. Silently omitted (not "n/a") when unavailable
    (older results.json, or a curve too short to have residual degrees of
    freedom) - the point estimate alone still stands on its own.
    """
    if ci95 is None:
        return ""
    lo, hi = ci95
    return f" (95% CI: {lo:.4g}–{hi:.4g} /hr)"


def _decay_results_table_cell_values(results, settings):
    """(row, col) -> replacement text for every cell in the decay
    results-table template that needs a real number - decay-mode
    counterpart to _results_table_cell_values. Row indices (0-based) match
    templates/decay_results_table_template.docx exactly:

      1 Average fluence rate           7  Effective ACHeff CFD measured
      2 Susceptibility constant Z      8  ACH efficiency
      3 (fixed IC - always "1")        9  eACHeff CFD measured
      4 Calculated eACH (well-mixed)  10  Total Room ACHeff+eACHeff (measured)
      6 ACH (air) measured            12  Total pathogen reduction/hour
                                      13  Room pathogen removal efficacy
                                      15  True UVGI Effectiveness

    (Row 14, "Total ACH in room (ACH+eACH_uv)" - the nominal/well-mixed
    total - was removed from the template: it only duplicated numbers
    already shown elsewhere (ACH + row 4) and risked being read as another
    measured figure alongside row 10.)

    "ACH (air) measured" (row 6) is the flow-rate/BC-based delivery check
    (setup_case's check_ach_delivery) - how much air is actually flowing.
    "Effective ACHeff CFD measured" (row 7) is the decay-curve-fit
    ventilation rate from the UV-off control run - how fast pathogen is
    actually removed, which can differ from row 6 due to imperfect mixing
    (confirmed directly on a real case: ~99% air delivery but only ~73%
    contaminant-removal-equivalent ACH). "ACH efficiency" (row 8) isolates
    that mixing-only gap: row7/row6.

    "Total average Pathogen reduction" (row 12) and "True UVGI
    Effectiveness" (row 15) are the same analytical steady-state figure
    (see _decay_reduction_ratio) - eACHeff/(ACHeff+eACHeff), using the
    measured ACHeff/eACHeff pair. "Room pathogen removal efficacy"
    (row 13) is a different ratio - measured total removal rate over the
    nominal/well-mixed total (ventilation_ach + eACH_uv_well_mixed) -
    computed internally even though that nominal total no longer gets its
    own row.
    """
    values = {}
    if results.get("fluence_mean") is not None:
        values[(1, 1)] = f"{results['fluence_mean']:.4g} µW/cm²"
    if settings.get("z-value") is not None:
        values[(2, 1)] = str(settings["z-value"])
    values[(3, 1)] = "1"  # decay mode's fixed initial condition (write_initial_fields default)

    eACH_well_mixed = results.get("eACH_uv_well_mixed")
    if eACH_well_mixed is not None:
        values[(4, 1)] = f"{eACH_well_mixed:.4g} /hr"

    ach_air_measured = (results.get("ach_delivery") or {}).get("measured_ach")
    if ach_air_measured is not None:
        values[(6, 1)] = f"{ach_air_measured:.4g} /hr"

    ach_eff = results.get("ventilation_ach_measured")
    if ach_eff is not None:
        values[(7, 1)] = f"{ach_eff:.4g} /hr{_ci_suffix(results.get('ventilation_ach_measured_ci95'))}"
    if ach_eff is not None and ach_air_measured:
        values[(8, 1)] = f"{ach_eff / ach_air_measured * 100:.1f}%"

    eACH_eff = results.get("eACH_uv_actual", results.get("eACH_uv_assuming_well_mixed"))
    eACH_eff_ci = results.get("eACH_uv_actual_ci95", results.get("eACH_uv_assuming_well_mixed_ci95"))
    if eACH_eff is not None:
        values[(9, 1)] = f"{eACH_eff:.4g} /hr{_ci_suffix(eACH_eff_ci)}"

    total_measured = (ach_eff + eACH_eff) if (ach_eff is not None and eACH_eff is not None) else None
    if total_measured is not None:
        values[(10, 1)] = f"{total_measured:.4g} /hr"
        # "True" UVGI Effectiveness and "Total average Pathogen reduction"
        # are algebraically identical here (both reduce to eACH/(ACH+eACH))
        # - same convention steady-state's own results table already uses
        # for its "True"/"Simple" UVGI Effectiveness rows, rather than
        # inventing a divergent formula that isn't actually there.
        reduction = _decay_reduction_ratio(eACH_eff, ach_eff)
        if reduction is not None:
            values[(12, 1)] = f"{reduction * 100:.1f}%"
            values[(15, 1)] = f"{reduction * 100:.1f}%"

    # Nominal/well-mixed total (ACH + eACH_uv well-mixed) - used for the
    # efficacy ratio below, but no longer has its own row (see docstring).
    total_nominal = results.get("total_ach_well_mixed")
    if total_measured is not None and total_nominal:
        values[(13, 1)] = f"{total_measured / total_nominal * 100:.1f}%"

    return values


def _fill_decay_results_table(table, results, settings):
    """Fill the decay results-table template's cells with real numbers, in
    place - decay-mode counterpart to _fill_results_table."""
    for (row, col), text in _decay_results_table_cell_values(results, settings).items():
        _set_paragraph_text(table.rows[row].cells[col].paragraphs[0], text)


def _monitoring_point_level_stats(results, point_data):
    """(T_point/T_avg ratio, point's detrended CV), both from trailing-
    window means/residuals (decay_analysis.windowed_stats_detrended) -
    the LEVEL a point sits at relative to the room average, which the
    reduction%/eACH figures never capture (those measure decay RATE, and a
    point can track the room average's rate closely while still sitting
    at a persistently different absolute level - confirmed directly: a
    real monitoring point ran at ~0.58-0.63x the room average for most of
    a run while its own eACH came out within ~1% of the room's). This
    matters for anything downstream that cares about absolute
    concentration at a specific location (e.g. infection-probability
    estimates), not just how fast it falls.

    None, None if either curve is missing - callers should skip silently
    (older results.json, or a run that predates recording t_seconds).
    """
    room = results.get("decay_curve") or {}
    t_room, T_room = room.get("t_seconds"), room.get("volAverage_T")
    t_point, T_point = point_data.get("t_seconds"), point_data.get("volAverage_T")
    if not (t_room and T_room and t_point and T_point):
        return None, None
    room_mean, *_ = windowed_stats_detrended(t_room, T_room)
    point_mean, _, point_cv, *_ = windowed_stats_detrended(t_point, T_point)
    ratio = point_mean / room_mean if room_mean else None
    return ratio, point_cv


def _fill_decay_monitoring_table(table, results):
    """Fill the decay template's per-point monitoring table (fixed row
    labels, e.g. "Patient pathogen reduction") by matching each row's own
    label text against a monitoring point name, rather than assuming a
    fixed point order - a differently-named or missing point just leaves
    that row as the template's own "xx.x %" placeholder.

    Per-point reduction uses the SAME analytical steady-state ratio as the
    room-average row (_decay_reduction_ratio), combining the room's own
    measured ACHeff (no separate per-point UV-off control curve is
    computed - see compute_monitoring_results) with that point's own
    eACH_uv_assuming_well_mixed (nominal-ACH-corrected, the only per-point
    measurement available) - a reasonable approximation, slightly less
    rigorous than the room-average version above which benefits from the
    real measured-ACH correction. The T_point/T_avg ratio and detrended CV
    (_monitoring_point_level_stats) are appended alongside it, since
    reduction% alone only reports rate, not the absolute level a point
    sits at relative to the room average.
    """
    monitoring = results.get("monitoring") or {}
    ach_eff = results.get("ventilation_ach_measured", results.get("ventilation_ach"))
    if ach_eff is None:
        return
    for row in table.rows:
        label = row.cells[0].text
        match = next((data for name, data in monitoring.items() if name.lower() in label.lower()), None)
        if match is None or match.get("eACH_uv_assuming_well_mixed") is None:
            continue
        reduction = _decay_reduction_ratio(match["eACH_uv_assuming_well_mixed"], ach_eff)
        if reduction is None:
            continue
        text = f"{reduction * 100:.1f}%"
        ratio, cv_detrended = _monitoring_point_level_stats(results, match)
        if ratio is not None:
            text += f" (T/Tavg={ratio:.2f}"
            if cv_detrended is not None:
                text += f", CV={cv_detrended * 100:.1f}%"
            text += ")"
        _set_paragraph_text(row.cells[1].paragraphs[0], text)


def _ach_source_note(res):
    """Appended to Reduction/measured-ACH rows: whether these were derived
    from the extrapolated T-infinity (see decay_analysis.fit_asymptotic_value)
    or the plain windowed average - see run_steady_state_scenario's
    "ach_source" field. Absent (empty string) for older results.json files
    that predate this distinction.
    """
    source = res.get("ach_source")
    if source == "extrapolated_T_infinity":
        return " (using extrapolated T∞)"
    if source == "windowed_average":
        return " (using windowed average - T∞ extrapolation unavailable)"
    return ""


def _phase_ss_rows(phase_num, uv_note, phase):
    """Steady-state phase1/phase2 rows: a trailing-window moving average +
    CV (see decay_analysis.windowed_stats_detrended) when the live
    per-iteration data is present, falling back to the old plain-T_ss row
    (exact original wording, e.g. "Phase 1 T_ss (no UV)") for older
    results.json files that predate live tracking.

    Extrapolated T-infinity (decay_analysis.fit_asymptotic_value) is an
    independent, model-based estimate of the true n->infinity value - not
    just another window - shown as its own row only when the fit
    succeeded (None for older results.json files, or when the fit didn't
    converge/the data isn't well-described by a single exponential).
    """
    plateau_note = f"({'plateaued' if phase['converged'] else 'NOT fully plateaued'}, " \
                    f"{phase['iterations']} iterations)"
    span = phase.get("T_ss_window_span")
    if span is None:
        return [(f"Phase {phase_num} T_ss ({uv_note})", f"{phase['T_ss']:.4g} {plateau_note}")]
    cv = phase.get("T_ss_cv")
    rows = [
        (f"Phase {phase_num} moving average ({uv_note}, last {span:.4g} iterations)",
         f"{phase['T_ss']:.4g} {plateau_note}"),
        (f"Phase {phase_num} CV ({uv_note}, last {span:.4g} iterations)",
         f"{cv * 100:.1f}%" if cv is not None else "n/a"),
    ]
    Tinf = phase.get("T_inf_extrapolated")
    if Tinf is not None:
        detail = phase.get("T_inf_extrapolation_detail") or {}
        fit_cv = detail.get("fit_cv")
        fit_cv_text = f", fit CV={fit_cv * 100:.2f}%" if fit_cv is not None else ""
        rows.append((f"Phase {phase_num} extrapolated T∞ ({uv_note}, n→∞)",
                      f"{Tinf:.4g}{fit_cv_text}"))
    return rows

def _total_ach_row(results):
    ach = results.get("ventilation_ach_measured")
    eACH_uv = results.get("eACH_uv_steady_state_corrected")
    if ach is None or eACH_uv is None:
        return "n/a"
    return f"{ach + eACH_uv:.4g} /hr"


# The approved, hand-designed steady-state "Results" table format (see
# CHANGELOG) - a fixed 24-row/2-col table with real Word footnotes
# explaining each derived quantity's equation. The row layout/wording/
# footnotes themselves live in the bundled template docx
# (templates/results_table_template.docx, with "<...>" placeholders for
# this module to fill); this module only ever fills numbers into it, never
# restructures it.
_RESULTS_TABLE_TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "results_table_template.docx"


def _results_table_cell_values(results, settings):
    """(row, col) -> replacement text for every cell in the results-table
    template that needs a real number. Cells not in this dict are left as
    the template shipped them (section headers, static labels). Missing
    source data (older results.json, or a phase whose T-infinity fit never
    converged) degrades to an explicit "n/a" rather than a stale/blank cell.
    """
    p1, p2 = results.get("phase1") or {}, results.get("phase2") or {}
    ach, z = settings.get("ach"), settings.get("z-value")
    values = {}

    if results.get("fluence_mean") is not None:
        values[(1, 1)] = f"{results['fluence_mean']:.4g} µW/cm²"
    if z is not None:
        values[(2, 1)] = str(z)
    values[(3, 1)] = f"{results.get('target_T_ss', '?')}"
    if results.get("injection_rate_total") is not None:
        values[(4, 1)] = f"{results['injection_rate_total']:.4g} "
    if results.get("eACH_uv_well_mixed") is not None:
        values[(5, 1)] = f"{results['eACH_uv_well_mixed']:.4g} /hr"

    # Labels are always rewritten (never left as the template's own "<...>"
    # placeholder), falling back to a window-size-agnostic phrase for
    # older results.json files that predate T_ss_window_frac/_span.
    frac1 = p1.get("T_ss_window_frac")
    frac1_desc = f"{frac1 * 100:.0f}% of last results" if frac1 is not None else "a trailing window of results"
    values[(7, 0)] = f"Steady state T, calculated from  moving average ({frac1_desc})"
    values[(8, 0)] = f"CV ({frac1_desc})"
    if p1.get("T_ss") is not None:
        plateau = "plateaued" if p1.get("converged") else "NOT plateaued"
        values[(7, 1)] = f"{p1['T_ss']:.4g} ({plateau})"
    if p1.get("T_ss_cv") is not None:
        values[(8, 1)] = f"{p1['T_ss_cv'] * 100:.1f}%"
    if p1.get("T_inf_extrapolated") is not None:
        fit_cv = (p1.get("T_inf_extrapolation_detail") or {}).get("fit_cv")
        fit_cv_text = f", fit CV={fit_cv * 100:.2f}%" if fit_cv is not None else ""
        values[(9, 1)] = f"{p1['T_inf_extrapolated']:.4g}{fit_cv_text}"
    else:
        values[(9, 1)] = "n/a (extrapolation unavailable)"

    span2 = p2.get("T_ss_window_span")
    window2_desc = f"last {int(span2)} iterations" if span2 is not None else "a trailing window"
    values[(11, 0)] = f"Steady State TSS2, calculated from moving average ({window2_desc})"
    values[(12, 0)] = f"CV ({window2_desc})"
    if p2.get("T_ss") is not None:
        values[(11, 1)] = f"{p2['T_ss']:.4g} "
    if p2.get("T_ss_cv") is not None:
        values[(12, 1)] = f"{p2['T_ss_cv'] * 100:.1f}%"
    if p2.get("T_inf_extrapolated") is not None:
        fit_cv = (p2.get("T_inf_extrapolation_detail") or {}).get("fit_cv")
        fit_cv_text = f", fit CV={fit_cv * 100:.2f}%" if fit_cv is not None else ""
        values[(13, 1)] = f"{p2['T_inf_extrapolated']:.4g}{fit_cv_text}"
    else:
        values[(13, 1)] = "n/a (extrapolation unavailable)"

    # Prefer reduction_pct_corrected (uses a T_ss1 implied by a real UV-off
    # control run's measured ventilation rate - see
    # compute_corrected_eACH_uv_from_control's docstring) over plain
    # reduction_pct (Phase 1's own point-source buildup, which can be
    # biased low by mixing-transport lag before it's fully converged) -
    # falls back to the uncorrected field for older results/single runs
    # that never had a control run.
    reduction_pct = results.get("reduction_pct_corrected", results.get("reduction_pct"))
    if reduction_pct is not None:
        values[(15, 1)] = f"{reduction_pct:.1f}% "
    ach_measured = results.get("ventilation_ach_measured")
    if ach_measured is not None:
        values[(16, 0)] = f"Effective pathogen (mechanical) ACHeff{_ach_source_note(results)}"
        values[(16, 1)] = f"{ach_measured:.4g} /hr "
    if ach_measured is not None and ach:
        values[(17, 1)] = f"{ach_measured / ach * 100:.1f}%"
    if results.get("eACH_uv_steady_state_corrected") is not None:
        values[(18, 1)] = f"{results['eACH_uv_steady_state_corrected']:.4g} /hr "
    if results.get("eACH_uv_steady_state") is not None:
        values[(19, 1)] = f"{results['eACH_uv_steady_state']:.4g} /hr"
    values[(20, 1)] = _total_ach_row(results)
    if reduction_pct is not None:
        # "True" and "Simple" UVGI effectiveness are algebraically
        # identical regardless of which ACH basis feeds the eACH_uv/ACHeff
        # pair (both always reduce to 1 - T_ss2/T_ss1, whichever T_ss1 that
        # is) - approved as showing the same number in both rows rather
        # than inventing a divergent "True" formula that isn't actually
        # there.
        values[(22, 1)] = f"{reduction_pct:.1f}%"
        values[(23, 1)] = f"{reduction_pct:.1f}%"
    return values


def _results_table_footnote_texts(results):
    """Footnote id -> replacement text, for the footnotes whose "<...>"
    equation/point-count placeholders this module fills. Only ids 1, 2, 4,
    5, 6, 7, 9, 10 are touched - 3, 8, 11, 12 were already complete in the
    approved template and are left exactly as shipped.
    """
    p1_n = len((results.get("phase1") or {}).get("live", {}).get("t", []))
    p2_n = len((results.get("phase2") or {}).get("live", {}).get("t", []))
    fit1_n = p1_n - int(p1_n * 0.5)
    fit2_n = p2_n - int(p2_n * 0.5)
    return {
        "1": (" Calculated based on given ACH and target well mixed steady state T with equation "
              "Tinj = V × (ACH/3600) × target_T_ss. If more than one injection source is "
              "given, the amount is equally divided by the quantity of sources"),
        "2": " Assuming well mixed condition eACH = Z*Eavg × 3.6 (unit conversion factor)",
        "4": (f" Extrapolated based on {fit1_n} simulation points (trailing 50% of Phase 1's "
              "iteration history) with equation TSS1∞ = TSS1 − A·exp(−n/τ), "
              "fit by nonlinear regression. T is the room average"),
        "5": (f" Extrapolated based on {fit2_n} simulation points (trailing 50% of Phase 2's "
              "iteration history) with equation TSS2∞ = TSS2 − A·exp(−n/τ), "
              "fit by nonlinear regression. T is the room average"),
        "6": (" Based on average T and sum of mechanical ACH and UV based eACH Using extrapolated "
              "T∞, calculated with equation reduction = (1 − TSS2∞/TSS1∞) × 100%"),
        "7": (" In not well mixed conditions the actual pathogen removal will be different than the "
              "mechanical air exchange rate (ACH). This value ACHeff is calculated from the simulated "
              "T values over time of phase 1 (no UV) by approximating the data of T(t) to equation "
              "ACHeff = (Tinj / (V × TSS1∞)) × 3600"),
        "9": (" Based on true mechanical ventilation ACHeff (using extrapolated T∞) calculated "
              "with equation eACHCFD = ACHeff × (TSS1∞/TSS2∞ − 1)"),
        "10": (" Based on rated ventilation ACH (and using extrapolated T∞) calculated with "
               "equation eACHCFD_s = eACHCFD / EACHeff"),
    }


def _set_paragraph_text(paragraph, new_text):
    """Replace a paragraph's visible text with new_text, leaving any
    footnoteReference run (and its exact position/formatting) untouched.
    """
    runs = list(paragraph.runs)
    text_runs = [run for run in runs if run._r.find(qn("w:footnoteReference")) is None]
    if not text_runs:
        return
    text_runs[0].text = new_text
    for run in text_runs[1:]:
        run.text = ""


def _fill_results_table(table, results, settings):
    """Fill the approved results-table template's table cells with real
    numbers, in place. Footnote text (word/footnotes.xml) can't be edited
    this way - python-docx doesn't parse that part into a live XML tree
    (see _patch_results_table_footnotes, applied post-save instead).

    Takes the actual Table object (captured by the caller right after
    loading the template, while it was still the document's only table) -
    not doc.tables[0], which re-resolves by body position and would
    silently pick up whatever table now sits first after other sections
    (metadata, Room Setup, ...) get relocated in front of it.
    """
    for (row, col), text in _results_table_cell_values(results, settings).items():
        _set_paragraph_text(table.rows[row].cells[col].paragraphs[0], text)


def _patch_results_table_footnotes(doc_out_path, results):
    """Fill the "<...>" equation/count placeholders in the results table's
    real Word footnotes. Must run on an already-saved .docx (python-docx
    only exposes footnotes.xml as opaque bytes, not a live XML tree - see
    _fill_results_table), so this reopens the file as a zip, patches just
    that one part, and rewrites the zip.

    Marker-run trap: a footnote's own auto-number (w:footnoteRef, styled
    small/superscript via the FootnoteReference character style) is a
    *different* element than w:footnoteReference (the body-side pointer
    into a footnote) - excluding only the latter treats the number marker
    itself as "the first text run" and its replacement text inherits the
    marker's small superscript style. Both must be excluded.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": W}
    footnote_values = _results_table_footnote_texts(results)

    with zipfile.ZipFile(doc_out_path) as zin:
        contents = {name: zin.read(name) for name in zin.namelist()}

    footnotes_xml = etree.fromstring(contents["word/footnotes.xml"])
    for footnote in footnotes_xml.findall("w:footnote", ns):
        fid = footnote.get(f"{{{W}}}id")
        if fid not in footnote_values:
            continue
        runs = footnote.findall(".//w:r", ns)
        text_runs = [
            r for r in runs
            if r.find(f"{{{W}}}footnoteRef") is None and r.find(f"{{{W}}}footnoteReference") is None
        ]
        if not text_runs:
            continue
        t = text_runs[0].find(f"{{{W}}}t")
        if t is None:
            t = etree.SubElement(text_runs[0], f"{{{W}}}t")
        t.text = footnote_values[fid]
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        for r in text_runs[1:]:
            tt = r.find(f"{{{W}}}t")
            if tt is not None:
                tt.text = ""
    contents["word/footnotes.xml"] = etree.tostring(
        footnotes_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(doc_out_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)


def _relocate_after(doc, after_element, before_element):
    """Move every body element currently AFTER after_element to just
    before before_element instead, preserving relative order (skips the
    trailing w:sectPr, which must stay last). Content added via the normal
    python-docx API always lands at the end of the body - when the
    document was seeded from the results-table template (so its Results
    heading/note/table start out already present, at the very top), this
    moves everything built afterwards (title, room setup, case setup) back
    in front of that template content instead, restoring the usual
    top-to-bottom report order.
    """
    body = doc.element.body
    to_move, found = [], False
    for child in list(body):
        if found:
            if child.tag == qn("w:sectPr"):
                break
            to_move.append(child)
        elif child is after_element:
            found = True
    for element in to_move:
        before_element.addprevious(element)


def _monitoring_rows(monitoring):
    """Row list for monitoring locations, if any were computed. Handles both
    decay's shape ({name: {t_seconds, volAverage_T, eACH_uv_assuming_well_mixed?}})
    and steady-state's shape ({name: {phase1: {...}, phase2: {...}}}).
    """
    if not monitoring:
        return []
    rows = []
    for name, data in monitoring.items():
        if "phase1" in data:
            p1, p2 = data["phase1"], data["phase2"]
            T1, T2, reduction_pct, basis = point_reduction_basis(p1, p2)
            value = f"T_ss1={T1:.4g}, T_ss2={T2:.4g}" if T1 is not None and T2 is not None else "n/a"
            if basis == "extrapolated_T_infinity":
                value += " (extrapolated)"
            cv1, cv2 = p1.get("T_ss_cv"), p2.get("T_ss_cv")
            if cv1 is not None or cv2 is not None:
                cv1_text = f"{cv1 * 100:.1f}%" if cv1 is not None else "n/a"
                cv2_text = f"{cv2 * 100:.1f}%" if cv2 is not None else "n/a"
                value += f" (CV1={cv1_text}, CV2={cv2_text})"
            if reduction_pct is not None:
                value += f", reduction={reduction_pct:.1f}%"
        else:
            T_final = data["volAverage_T"][-1] if data["volAverage_T"] else None
            value = f"final volAverage(T)={T_final:.4g}" if T_final is not None else "n/a"
            if data.get("eACH_uv_assuming_well_mixed") is not None:
                value += f", eACH_uv={data['eACH_uv_assuming_well_mixed']:.4g}/hr"
        rows.append((name, value))
    return rows


def _room_setup_rows(room, settings):
    """_ROW_LABELS_ROOM's base rows, plus a 2nd inlet/outlet row (only when
    enabled), an injection-point row (steady-state runs only), and one row
    per monitoring point - all read straight from run_settings.json, same
    provenance as everything else in this table.
    """
    rows = [(label, fn(room, settings)) for label, fn in _ROW_LABELS_ROOM]
    if settings.get("inlet2-enable"):
        rows.append(("Inlet 2", f"{settings['inlet2-wall']}, y={settings['inlet2-y-input']:.3g}m "
                                 f"z={settings['inlet2-z-input']:.3g}m, size={settings['inlet2-size-w']:.3g}x"
                                 f"{settings['inlet2-size-h']:.3g}m"))
    if settings.get("outlet2-enable"):
        rows.append(("Outlet 2", f"{settings['outlet2-wall']}, y={settings['outlet2-y-input']:.3g}m "
                                  f"z={settings['outlet2-z-input']:.3g}m, size={settings['outlet2-size-w']:.3g}x"
                                  f"{settings['outlet2-size-h']:.3g}m"))
    source_center = settings.get("source_center")
    if source_center and all(v is not None for v in source_center):
        rows.append(("Injection point",
                      f"({source_center[0]:.3g}, {source_center[1]:.3g}, {source_center[2]:.3g}) m"))
    for pt in settings.get("monitoring_points") or []:
        rows.append((f"Monitoring point: {pt['name']}",
                      f"({pt['x']:.3g}, {pt['y']:.3g}, {pt['z']:.3g}) m, "
                      f"box={pt['cells_per_side']} cells/side"))
    return rows


def _trust_status_rows(results):
    """Plain-language rows summarizing whether this run's flow field,
    ventilation delivery, and (for steady-state runs) Phase 1 mass balance
    are trustworthy - kept as separate rows, not one blended verdict,
    because they check genuinely different things and can disagree: a
    perfectly-converged-looking contaminant curve can still be riding on a
    flow field that never itself converged (only accepted via bounded
    oscillation), and neither of those says whether the mesh/BCs are even
    delivering the intended ventilation rate at all - confirmed directly
    on a real case where all of flow-convergence and T-plateau looked fine
    while the inlet was silently delivering only ~38% of its target flow.
    Missing/None fields (older result files, or flow convergence skipped)
    are reported as "not available" rather than guessed at.
    """
    rows = []

    flow_converged = results.get("flow_converged")
    if flow_converged is None:
        rows.append(("Flow field", "Not available (older run, or flow convergence was skipped)"))
    elif flow_converged:
        rows.append(("Flow field", "Converged (flow-convergence tolerance satisfied)"))
    else:
        rows.append(("Flow field", "Accepted via bounded oscillation - never fully converged "
                                    "(common for a jet/fan impinging on a wall or floor; see "
                                    "Advanced Settings > Flow oscillation acceptance)"))

    ach = results.get("ach_delivery")
    if ach is None:
        rows.append(("Ventilation delivery", "Not checked (older run)"))
    else:
        verdict = "OK" if ach["within_tolerance"] else "MISMATCH - do not trust downstream results"
        rows.append(("Ventilation delivery", f"{ach['measured_ach']:.3g}/hr measured vs "
                                              f"{ach['nominal_ach']:.3g}/hr nominal "
                                              f"({ach['ratio']:.0%}) - {verdict}"))

    mass_balance = results.get("phase1", {}).get("mass_balance") if "phase1" in results else None
    if mass_balance is not None:
        verdict = "balanced" if mass_balance["within_tolerance"] else "NOT balanced - Phase 1 may not be converged"
        rows.append(("Phase 1 mass balance", f"outlet removal / injection = {mass_balance['ratio']:.0%} "
                                              f"- {verdict}"))

    return rows


def _add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    return table


def generate_report_docx(case_dir, out_path):
    """Build the report from run_settings.json + results.json in case_dir.
    Raises FileNotFoundError with a clear message if either is missing (no
    completed run to report on yet).
    """
    settings_path = Path(case_dir) / "run_settings.json"
    results_path = Path(case_dir) / "results.json"
    if not settings_path.exists() or not results_path.exists():
        raise FileNotFoundError(
            f"{case_dir} doesn't have both run_settings.json and results.json - "
            f"run a full simulation to completion first."
        )
    with open(settings_path) as f:
        settings = json.load(f)
    with open(results_path) as f:
        # migrate_result_keys: results.json files written before the
        # 2026-09-02 rename use eACH_uv_effective/_corrected. Reports must
        # keep opening them.
        results = migrate_result_keys(json.load(f))

    guv_path = settings.get("guv_path")
    if not guv_path:
        raise FileNotFoundError(
            f"{case_dir}/run_settings.json has no recorded project path - "
            f"this case directory predates report support; rerun a full "
            f"simulation here to enable report generation."
        )
    project = Project.load(guv_path)
    room = next(iter(project.rooms.values()))

    # injection_rate_total predates this field for case dirs from before it
    # was added to steady_state_pipeline.py - it's a deterministic function
    # of room volume/ACH/target_T_ss (see compute_source_strength), so an
    # older results.json missing it can still show the real number instead
    # of "n/a" rather than needing a rerun.
    if "phase1" in results and results.get("injection_rate_total") is None \
            and results.get("target_T_ss") and settings.get("ach"):
        results = dict(results)
        results["injection_rate_total"] = compute_source_strength(
            room.x * room.y * room.z, settings["ach"], results["target_T_ss"])

    # eACH_uv_well_mixed predates this field for steady-state case dirs from
    # before app.py's _run_steady_state started copying it out of setup_case()'s
    # summary - it's exactly Z * fluence_mean * 3.6 (see fluence.py's
    # compute_inactivation_rate/compute_well_mixed_eACH: linear in E, so
    # order of averaging doesn't matter), and fluence_mean has always been
    # saved, so an older results.json can still show the real number.
    if "phase1" in results and results.get("eACH_uv_well_mixed") is None \
            and results.get("fluence_mean") is not None and settings.get("z-value") is not None:
        results = dict(results)
        results["eACH_uv_well_mixed"] = settings["z-value"] * results["fluence_mean"] * 3.6

    fan_kwargs = {}
    if settings.get("fan-enable"):
        direction = (0, 0, -1) if settings["fan-direction"] == "down" else (0, 0, 1)
        fan_kwargs = dict(
            fan_speed=settings["fan-speed"], fan_disk_radius=settings["fan-radius"],
            fan_disk_thickness=settings["fan-thickness"],
            fan_center=(settings["fan-x-input"], settings["fan-y-input"], settings["fan-z-input"]),
            fan_direction=direction,
        )

    opening2_kwargs = {}
    if settings.get("inlet2-enable"):
        opening2_kwargs.update(
            inlet2_wall=settings["inlet2-wall"],
            inlet2_center=center_frac_for_wall(
                settings["inlet2-wall"], settings["inlet2-y-input"], settings["inlet2-z-input"], room),
            inlet2_size=(settings["inlet2-size-w"], settings["inlet2-size-h"]),
        )
    if settings.get("outlet2-enable"):
        opening2_kwargs.update(
            outlet2_wall=settings["outlet2-wall"],
            outlet2_center=center_frac_for_wall(
                settings["outlet2-wall"], settings["outlet2-y-input"], settings["outlet2-z-input"], room),
            outlet2_size=(settings["outlet2-size-w"], settings["outlet2-size-h"]),
        )

    source_center = settings.get("source_center")
    injection_center = (tuple(source_center)
                         if source_center and all(v is not None for v in source_center) else None)

    fig = plot_case(
        room,
        inlet_wall=settings["inlet-wall"],
        inlet_center=(settings["inlet-y-input"] / room.y, settings["inlet-z-input"] / room.z),
        inlet_size=(settings["inlet-size-w"], settings["inlet-size-h"]),
        outlet_wall=settings["outlet-wall"],
        outlet_center=(settings["outlet-y-input"] / room.y, settings["outlet-z-input"] / room.z),
        outlet_size=(settings["outlet-size-w"], settings["outlet-size-h"]),
        injection_center=injection_center,
        monitoring_points=settings.get("monitoring_points"),
        title="", **fan_kwargs, **opening2_kwargs,
    )
    fig.update_layout(margin=dict(l=0, r=0, t=10, b=0), width=900, height=650)

    # Result curve (phase timeline for steady-state, decay curve for decay)
    # - only when the run actually recorded curve data, so older/minimal
    # case dirs still generate a report, just without this picture.
    if "phase1" in results:
        has_curve = bool(results["phase1"].get("decay_curve", {}).get("t")
                          and results["phase2"].get("decay_curve", {}).get("t"))
        curve_fig = steady_state_figure(results) if has_curve else None
    else:
        has_curve = bool(results.get("decay_curve", {}).get("t_seconds")
                          and results.get("eACH_uv_well_mixed") is not None)
        curve_fig = decay_figure(results) if has_curve else None
    if curve_fig is not None:
        curve_fig.update_layout(margin=dict(l=50, r=20, t=30, b=45), width=900, height=500)

    # Rendered pictures are staged in a real temp directory, not next to
    # out_path - if write_image() or doc.save() blows up partway through
    # (kaleido is known to be flaky - see test_report.py), nothing gets left
    # behind that could be mistaken for the report itself.
    with tempfile.TemporaryDirectory() as tmp_dir:
        image_path = Path(tmp_dir) / "preview.png"
        fig.write_image(str(image_path))
        curve_image_path = None
        if curve_fig is not None:
            curve_image_path = Path(tmp_dir) / "curve.png"
            curve_fig.write_image(str(curve_image_path))
        _write_report_docx(doc_out_path=out_path, case_dir=case_dir, guv_path=guv_path,
                            settings=settings, results=results, room=room,
                            image_path=image_path, curve_image_path=curve_image_path)
    return out_path


def _write_report_docx(doc_out_path, case_dir, guv_path, settings, results, room,
                        image_path, curve_image_path):
    # Both report kinds are seeded from an approved template (real Word
    # footnotes for steady-state; a fixed results+monitoring table layout
    # for decay - see _RESULTS_TABLE_TEMPLATE_PATH/_DECAY_RESULTS_TABLE_
    # TEMPLATE_PATH) instead of a blank Document() - the template's own
    # content starts out at the very TOP of the body, so everything built
    # below gets relocated in front of it just before saving
    # (_relocate_after) to restore normal top-to-bottom report order.
    is_steady_state = "phase1" in results
    doc = Document(str(_RESULTS_TABLE_TEMPLATE_PATH if is_steady_state else _DECAY_RESULTS_TABLE_TEMPLATE_PATH))
    results_anchor = doc.paragraphs[0]._p
    # Captured now, while the template's own table(s) are still the only
    # ones in the document - doc.tables[0] would re-resolve by body
    # position later and silently pick up a different table once other
    # sections get relocated in front of it (see _fill_results_table).
    results_table = doc.tables[0]
    monitoring_table = doc.tables[1] if not is_steady_state and len(doc.tables) > 1 else None
    # The decay template's own mixing-uniformity note paragraph - found by
    # its own content, NOT doc.paragraphs[-1] (the template has a trailing
    # *empty* paragraph after it, which silently ate the replacement text
    # via _set_paragraph_text's "no text runs -> no-op" guard - confirmed
    # directly: the shipped example text, "...(Phase 1)...(Phase 2)...",
    # survived untouched in a generated report). Captured as a Paragraph
    # object (not an index) so it stays correctly identified even after
    # _relocate_after reorders everything in front of it.
    decay_uniformity_para = None
    if not is_steady_state:
        decay_uniformity_para = next(p for p in doc.paragraphs if "well mixed" in p.text.lower())
    # after_element for _relocate_after: the template's own last piece of
    # content - the results table for steady-state (nothing follows it),
    # or the true trailing paragraph for decay (its own blank paragraph,
    # not the uniformity note - there's more after the note itself).
    template_tail = results_table._tbl if is_steady_state else doc.paragraphs[-1]._p

    doc.add_heading("GUV-CFD Simulation Report", level=1)
    doc.add_paragraph(f"Illuminate room design file: {guv_path}")
    doc.add_paragraph(f"CFD Project file: {settings.get('settings_path') or 'n/a'}")
    doc.add_paragraph(f"OpenFoam directory: {case_dir}")

    started_at, elapsed_seconds = _run_timing(case_dir, results)
    system_info = get_system_info()
    metadata_rows = []
    if started_at is not None:
        metadata_rows.append(("Simulation date", started_at.strftime("%Y-%m-%d %H:%M")))
    if elapsed_seconds is not None:
        metadata_rows.append(("Total elapsed time", _format_elapsed(elapsed_seconds)))
    metadata_rows.append(("CPU", system_info["cpu"]))
    if system_info["ram_gb"] is not None:
        metadata_rows.append(("RAM", f"{system_info['ram_gb']:.1f} GB"))
    if system_info["gpu"]:
        metadata_rows.append(("GPU", f"{system_info['gpu']} (not used - this simulation's "
                                      "OpenFOAM solve is CPU-only)"))
    _add_kv_table(doc, metadata_rows)

    doc.add_heading("Room Setup", level=2)
    _add_kv_table(doc, _room_setup_rows(room, settings))
    if settings.get("fan-enable"):
        _add_kv_table(doc, [(label, fn(settings)) for label, fn in _ROW_LABELS_FAN])

    doc.add_heading("Case Setup", level=2)
    doc.add_picture(str(image_path), width=Inches(6.0))

    trust_rows = _trust_status_rows(results)
    if trust_rows:
        doc.add_heading("Convergence & Trust", level=2)
        _add_kv_table(doc, trust_rows)

    # Relocate everything built so far (title/metadata, Room Setup, Case
    # Setup) to in front of the template's own content - see this
    # function's opening comment.
    _relocate_after(doc, template_tail, results_anchor)

    if is_steady_state:
        _fill_results_table(results_table, results, settings)
        if curve_image_path is not None:
            doc.add_picture(str(curve_image_path), width=Inches(6.0))
        if results.get("ventilation_ach_measured") is not None:
            doc.add_paragraph().add_run(_effective_ach_note(results)).italic = True

        monitoring_rows = _monitoring_rows(results.get("monitoring"))
        if monitoring_rows:
            doc.add_heading("Monitoring Results", level=2)
            _add_kv_table(doc, monitoring_rows)

        uniformity_note = mixing_uniformity_note(results)
        if uniformity_note:
            doc.add_paragraph().add_run(uniformity_note).italic = True
    else:
        _fill_decay_results_table(results_table, results, settings)
        if monitoring_table is not None:
            _fill_decay_monitoring_table(monitoring_table, results)
            _set_paragraph_text(decay_uniformity_para, mixing_uniformity_note(results) or "")
        if curve_image_path is not None:
            doc.add_picture(str(curve_image_path), width=Inches(6.0))
        if results.get("ventilation_ach_measured") is not None:
            # Decay mode has only ever had one way to measure this (a
            # dedicated UV-off control run) - unconditionally the control
            # wording, unlike steady-state's _effective_ach_note above.
            doc.add_paragraph().add_run(EFFECTIVE_ACH_NOTE_CONTROL).italic = True

    doc.save(doc_out_path)
    if is_steady_state:
        # word/footnotes.xml isn't reachable as a live XML tree through
        # python-docx (see _patch_results_table_footnotes) - has to be
        # patched as a post-save step against the file that was just
        # written, not the in-memory Document.
        _patch_results_table_footnotes(doc_out_path, results)
